# -*- coding: utf-8 -*-
"""Fills all missing sections into gadash_report.docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_PATH = "gadash_report.docx"

doc = Document(DOC_PATH)


# ── helpers ────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x2f, 0x4e)
    return p


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(text)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Navy background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A2F4E")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Column widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 (existing) — ספריות צד שלישי  (complete the empty header)
# ══════════════════════════════════════════════════════════════════════════════

add_para(doc, "להלן רשימת כל ספריות הצד השלישי בהן השתמשנו, מטרתן ומקורן:")

libs = [
    ("Flask 2.3.3",              "תשתית",   "מסגרת Web מינימליסטית — ניתוב, תבניות Jinja2, session"),
    ("gunicorn 21.2.0",          "תשתית",   "שרת WSGI לסביבת ייצור (Fly.io)"),
    ("python-telegram-bot 20.7", "תקשורת",  "ממשק אסינכרוני לבוט טלגרם — ConversationHandler, JobQueue"),
    ("gspread 6.1.2",            "נתונים",  "ממשק Python ל-Google Sheets API v4"),
    ("google-auth 2.29.0",       "נתונים",  "אימות Service Account מול Google APIs"),
    ("google-generativeai 0.8.5","AI",      "ממשק ל-Gemini 2.5 Flash לסיכומים חודשיים בעברית"),
    ("pandas 2.2.2",             "נתונים",  "מניפולציה וסינון DataFrame; ייצוא Excel/CSV"),
    ("numpy 1.26.4",             "נתונים",  "תלות של pandas — חישובים מספריים"),
    ("openpyxl 3.1.2",           "נתונים",  "כתיבת/קריאת קבצי .xlsx לייצוא/ייבוא"),
    ("Flask-WTF 1.2.1",          "אבטחה",   "ניהול CSRF tokens בטפסים"),
    ("Flask-Limiter 3.5.0",      "אבטחה",   "הגבלת קצב בקשות (rate limiting) על נתיבי התחברות"),
    ("Werkzeug 3.x",             "אבטחה",   "PBKDF2 password hashing לסיסמאות מנהל ועובדים"),
    ("python-dotenv 1.0.1",      "תשתית",   "טעינת משתני סביבה מקובץ .env בפיתוח מקומי"),
]

make_table(doc,
           ["ספרייה וגרסה", "תחום", "שימוש במערכת"],
           libs,
           col_widths=[5.5, 2.5, 8.0])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — צורת הכניסה למערכת
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "סעיף 8 — צורת הכניסה למערכת (Authentication)", level=1)

add_para(doc, "8.1 סוגי משתמשים והרשאות", bold=True)
add_para(doc, "במערכת שני תפקידים מובחנים, כל אחד עם session נפרד וסמכויות שונות:")

make_table(doc,
    ["תפקיד", "session key", "הרשאות"],
    [
        ("מנהל", 'session["logged_in"] = True',
         "גישה מלאה: צפייה, הוספה, עריכה, מחיקה, ייצוא, ייבוא, ניהול עובדים, לוג ביקורת"),
        ("עובד Web", 'session["worker_logged_in"] = True',
         "הזנת רשומות בלבד + צפייה בהיסטוריה האישית; אין גישה לנתוני אחרים"),
        ("עובד בוט", "Telegram user_id",
         "שיחה מודרכת להזנה; /undo לביטול אחרון; סטטיסטיקות אישיות"),
    ],
    col_widths=[3.0, 5.5, 8.0])

doc.add_paragraph()
add_para(doc, "8.2 שיטת האימות", bold=True)

auth_items = [
    "שם משתמש + סיסמה — המנהל מזין סיסמת מנהל; עובד Web מזין שם עובד וסיסמה",
    "סיסמת המנהל מאוחסנת כ-hash בגיליון Settings בפורמט Werkzeug PBKDF2-SHA256",
    "סיסמאות עובדים מאוחסנות כ-hash בגיליון Workers; תמיכה לאחור ב-SHA-256 ישן",
    "השוואה תמיד דרך check_password_hash() — סיסמה גולמית לא נשמרת",
    "ניהול המצב: Flask session חתומה עם SECRET_KEY, פג תוקף אחרי 8 שעות",
    "הגנה נוספת: rate limiting — לא יותר מ-5 ניסיונות שגויים בדקה לכל IP",
    "CSRF token בכל טופס POST — מאומת ב-before_request לפני כל כתיבה",
]
for item in auth_items:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc, "8.3 זרימת הכניסה", bold=True)
add_para(doc, "להלן תרשים זרימת תהליך הכניסה למערכת:")

flow_steps = [
    ("התחלה", "המשתמש פותח /login"),
    ("קלט", "בחירת תפקיד (מנהל / עובד) + הזנת שם וסיסמה"),
    ("בדיקת rate limit", "אם > 5 ניסיונות בדקה → הודעת שגיאה + חסימה"),
    ("אימות סיסמה", "check_password_hash() מול hash בגיליון Settings / Workers"),
    ("הצלחה", "session['logged_in']=True / session['worker_logged_in']=True → redirect"),
    ("כישלון", "_record_attempt() → הצגת מספר ניסיונות שנותרו"),
    ("session timeout", "אחרי 8 שעות ללא פעילות — ניתוב אוטומטי חזרה ל-/login"),
]
make_table(doc,
           ["שלב", "תיאור"],
           flow_steps,
           col_widths=[4.5, 12.0])

doc.add_paragraph()
add_para(doc, "8.4 Decorator אימות", bold=True)
add_para(doc, "כל נתיב מוגן עטוף ב-decorator המתאים:")

add_code_block(doc,
    "@login_required       # בודק session['logged_in']  → מנהל בלבד\n"
    "@worker_required      # בודק session['worker_logged_in'] → עובד בלבד\n\n"
    "# דוגמה:\n"
    "@app.route('/add', methods=['GET','POST'])\n"
    "@login_required\n"
    "def add(): ..."
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — עץ התיקיות
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "סעיף 9 — עץ התיקיות", level=1)
add_para(doc, "להלן מבנה קבצי הפרויקט. קבצי cache, __pycache__ וקבצי עזר זמניים הושמטו לשם בהירות:")

tree_text = (
    "gadash_bot/\n"
    "├── app.py                   # נקודת כניסה Flask — ניתובים, auth, CSRF, rate limit\n"
    "├── requirements.txt         # תלויות Python\n"
    "├── Dockerfile               # קונטיינר Python 3.10-slim\n"
    "├── fly.toml                 # הגדרות פריסה Fly.io — Amsterdam, 256MB\n"
    "├── .env.example             # דוגמה למשתני סביבה\n"
    "├── README.md                # הוראות התקנה והרצה\n"
    "│\n"
    "├── gadash/                  # חבילת הליבה\n"
    "│   ├── __init__.py\n"
    "│   ├── models.py            # WorkEntry dataclass + COLUMNS + VALID_TASKS\n"
    "│   ├── sheets.py            # כל I/O ל-Google Sheets + cache TTL 5 דקות\n"
    "│   ├── service.py           # create_entry() — שכבת שירות\n"
    "│   ├── bot.py               # בוט טלגרם — ConversationHandler (16 states)\n"
    "│   ├── audit.py             # לוג כפול: audit.log + גיליון AuditLog\n"
    "│   ├── workers.py           # CRUD עובדים + Werkzeug password hashing\n"
    "│   └── subscribers.py       # רשימת chat_ids לשידורים אוטומטיים\n"
    "│\n"
    "├── templates/               # תבניות Jinja2 (RTL, Bootstrap 5.3)\n"
    "│   ├── base.html            # תבנית בסיס — navbar, sidebar, loading spinner\n"
    "│   ├── base_worker.html     # תבנית בסיס לפורטל עובד\n"
    "│   ├── index.html           # לוח בקרה ראשי — טבלה, חיפוש, עריכה מוטבעת\n"
    "│   ├── worker_index.html    # פורטל עובד — טופס + היסטוריה\n"
    "│   ├── dashboard.html       # דשבורד ניתוח עם גרפי Chart.js\n"
    "│   ├── summary.html         # דוח חודשי pivot\n"
    "│   ├── field_report.html    # דוח שעות לפי חלקה\n"
    "│   ├── client_report.html   # דוח לקוח\n"
    "│   ├── fields_map.html      # מפת שדות — Leaflet.js\n"
    "│   ├── login.html           # מסך כניסה עם בורר תפקיד\n"
    "│   ├── add.html             # הוספת רשומה\n"
    "│   ├── edit.html            # עריכת רשומה\n"
    "│   ├── workers.html         # ניהול עובדים\n"
    "│   ├── audit.html           # לוג פעולות\n"
    "│   ├── import.html          # ייבוא Excel\n"
    "│   ├── 404.html             # עמוד שגיאה 404 בעברית\n"
    "│   └── 500.html             # עמוד שגיאה 500 בעברית\n"
    "│\n"
    "├── tests/\n"
    "│   └── test_app.py          # 57 בדיקות pytest\n"
    "│\n"
    "└── uml_*.puml               # קבצי PlantUML — use case, class, activity, state\n"
)
add_code_block(doc, tree_text)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — רשימת הקבצים
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "סעיף 10 — רשימת הקבצים", level=1)
add_para(doc, "טבלה המפרטת את הקבצים המרכזיים בפרויקט, שכבתם ותפקידם:")

file_rows = [
    ("app.py",                    "לוגיקה",  "נקודת כניסה Flask; כל הנתיבים, auth decorators, CSRF, rate limiting, error handlers"),
    ("gadash/models.py",          "לוגיקה",  "WorkEntry dataclass עם __post_init__ validation; COLUMNS (11 עמודות); VALID_TASKS"),
    ("gadash/sheets.py",          "נתונים",  "כל I/O ל-Google Sheets; cache TTL 5 דק' עם threading.Lock; CRUD: append, edit, delete, patch"),
    ("gadash/service.py",         "לוגיקה",  "create_entry() — שכבת שירות דקה המחברת בין ה-route לבין sheets.py ו-audit.py"),
    ("gadash/bot.py",             "לוגיקה",  "בוט טלגרם; ConversationHandler עם 16 מצבים; JobQueue לשידורים ב-08:00 ו-18:00"),
    ("gadash/audit.py",           "לוגיקה",  "לוג כפול: כתיבה מיידית ל-audit.log + flush אסינכרוני לגיליון AuditLog כל 30 שניות"),
    ("gadash/workers.py",         "נתונים",  "CRUD עובדים בגיליון Workers; _hash_pw() ב-Werkzeug; backward-compat ל-SHA-256 ישן"),
    ("gadash/subscribers.py",     "נתונים",  "ניהול רשימת chat_ids של מנויים לשידורי הבוט"),
    ("templates/base.html",       "מצג",     "תבנית בסיס RTL; Bootstrap 5.3 + navbar + sidebar offcanvas + loading overlay"),
    ("templates/index.html",      "מצג",     "לוח בקרה ראשי; טבלה עם עימוד; חיפוש גלובלי debounce; עריכה מוטבעת (dblclick/double-tap)"),
    ("templates/worker_index.html","מצג",    "פורטל עובד; טופס הזנה + autocomplete + כפתור undo-last"),
    ("templates/dashboard.html",  "מצג",     "דשבורד; Chart.js — daily trend, task pie, top clients, monthly bar"),
    ("templates/fields_map.html", "מצג",     "מפת שדות אינטראקטיבית — Leaflet.js + OpenStreetMap tiles"),
    ("Dockerfile",                "פריסה",   "FROM python:3.10-slim; COPY requirements.txt; pip install; COPY . .; gunicorn"),
    ("fly.toml",                  "פריסה",   "app=gadash-bot; region=ams; 256MB RAM; auto-stop; force HTTPS"),
    ("requirements.txt",          "תשתית",   "13 תלויות עם גרסאות קבועות; python-telegram-bot[job-queue] לתמיכה ב-JobQueue"),
    ("tests/test_app.py",         "בדיקות",  "57 בדיקות pytest; WorkEntry validation, Flask routes, security, dashboard cache, bot smoke tests"),
]

make_table(doc,
           ["קובץ", "שכבה", "תפקיד"],
           file_rows,
           col_widths=[4.5, 2.0, 10.0])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — סכימות הנתונים
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "סעיף 11 — סכימות הנתונים", level=1)

add_para(doc,
    "מקור הנתונים היחיד של המערכת הוא גיליון Google Sheets בשם \"Gadash Data\". "
    "הגיליון מכיל 6 worksheets המשמשים כ\"טבלאות\" של מסד הנתונים. "
    "אין בסיס נתונים רלציוני — הקשרים מיוצגים דרך ערכי מפתח משותפים.")

doc.add_paragraph()
add_para(doc, "11.1 תרשים ER (מושגי)", bold=True)
add_para(doc,
    "להלן תרשים ER מושגי המתאר את הקשרים בין הישויות. "
    "במקום Foreign Key רגיל, ערך השדה \"מזין\" ב-Sheet1 מתאים לשדה \"שם\" ב-Workers, "
    "ו\"שם חלקה\" ב-Sheet1 מתאים לשדה \"שם חלקה\" ב-FieldCoords.")

add_code_block(doc,
    "┌─────────────────┐        ┌──────────────────┐\n"
    "│   Sheet1        │        │   Workers        │\n"
    "│  (עבודות)       │        │  (עובדים)        │\n"
    "│─────────────────│        │──────────────────│\n"
    "│ שם לקוח         │        │ שם  ◄────────────┼── מזין (Sheet1)\n"
    "│ תאריך           │        │ password_hash    │\n"
    "│ עבודה           │        │ telegram_id      │\n"
    "│ שם חלקה ────────┼──────► │                  │\n"
    "│ גידול           │   ┌────┴──────────────────┘\n"
    "│ כמות            │   │\n"
    "│ שעות            │   ▼\n"
    "│ כלי             │  ┌────────────────────┐\n"
    "│ מפעיל           │  │   FieldCoords      │\n"
    "│ הערות           │  │────────────────────│\n"
    "│ מזין            │  │ שם חלקה (PK)       │\n"
    "└─────────────────┘  │ lat                │\n"
    "                     │ lng                │\n"
    "                     └────────────────────┘\n\n"
    "┌──────────────┐    ┌──────────────┐    ┌──────────────┐\n"
    "│  AuditLog    │    │  Settings    │    │  Subscribers │\n"
    "│──────────────│    │──────────────│    │──────────────│\n"
    "│ ts           │    │ key (PK)     │    │ chat_id      │\n"
    "│ action       │    │ value        │    └──────────────┘\n"
    "│ user         │    └──────────────┘\n"
    "│ detail       │\n"
    "└──────────────┘"
)

doc.add_paragraph()
add_para(doc, "11.2 פירוט הגיליונות", bold=True)

sheets_desc = [
    ("Sheet1", "נתוני עבודה ראשיים",
     "שורה 1 = כותרות; כל שורה = רשומת עבודה אחת. row_id = אינדקס DataFrame (0-based); שורת גיליון = row_id + 2"),
    ("Workers", "רישום עובדים",
     "שדה 'שם' הוא מפתח ראשי (טקסט). password_hash בפורמט Werkzeug PBKDF2. telegram_id מקשר לחשבון טלגרם"),
    ("Settings", "הגדרות מערכת (key-value)",
     "שורות: web_password (hash של סיסמת מנהל), worker_password. ערכים נטענים ב-startup ומוחלפים על env vars"),
    ("FieldCoords", "קואורדינטות GPS לחלקות",
     "שדה 'שם חלקה' = מפתח. lat/lng בטיפוס float. cache נפרד TTL 5 דקות. מוצג על מפת Leaflet"),
    ("AuditLog", "לוג ביקורת",
     "נכתב כל פעולת כתיבה. flush אסינכרוני מ-audit.log לגיליון כל 30 שניות. קיים גם כקובץ audit.log מקומי"),
    ("Subscribers", "מנויים לבוט",
     "chat_id של כל מי שהפעיל /start בבוט. משמש לשידורים האוטומטיים (08:00 יומי, 18:00 תזכורת)"),
]

add_heading(doc, "גיליונות הנתונים", level=2)
for sheet_name, title, desc in sheets_desc:
    add_para(doc, f"{sheet_name} — {title}", bold=True)
    add_para(doc, desc)
    doc.add_paragraph()

add_para(doc, "11.3 סכימת השדות — Sheet1 (נתוני עבודה)", bold=True)

sheet1_cols = [
    ("שם לקוח",  "text",    "חובה",  "שם הלקוח שעבורו בוצעה העבודה"),
    ("תאריך",    "text",    "חובה",  "פורמט YYYY-MM-DD; מאומת ב-WorkEntry.__post_init__()"),
    ("עבודה",    "enum",    "חובה",  "חריש / ריסוס / קציר / דיסוק / אחר (VALID_TASKS)"),
    ("שם חלקה", "text",    "רשות", "שם החלקה; FK ל-FieldCoords.שם חלקה"),
    ("גידול",    "text",    "רשות", "סוג הגידול (חיטה, תירס וכד')"),
    ("כמות",     "text",    "רשות", "כמות חופשית (30 דונם, 5 טון וכד')"),
    ("שעות",     "number",  "רשות", "שעות עבודה; מסוכם בדוחות"),
    ("כלי",      "text",    "רשות", "ציוד ששימש (טרקטור, ריסוסית וכד')"),
    ("מפעיל",    "text",    "רשות", "שם המפעיל"),
    ("הערות",    "text",    "רשות", "הערות חופשיות"),
    ("מזין",     "text",    "אוטו'","שם המשתמש שהזין ('Web', שם עובד, שם מלא מטלגרם)"),
]

make_table(doc,
           ["שדה", "סוג", "חובה?", "תיאור"],
           sheet1_cols,
           col_widths=[2.5, 2.0, 1.8, 10.0])

doc.add_paragraph()
add_para(doc, "11.4 סכימת גיליון Workers", bold=True)

workers_cols = [
    ("שם",            "text (PK)",  "שם העובד — ייחודי, משמש כמפתח ראשי"),
    ("password_hash", "text",       "Werkzeug PBKDF2-SHA256; backward-compat ל-SHA-256 ישן (64 תווי hex)"),
    ("telegram_id",   "text",       "chat_id של הטלגרם; ריק עד שהעובד מקשר חשבון דרך הבוט"),
]
make_table(doc,
           ["שדה", "סוג", "תיאור"],
           workers_cols,
           col_widths=[3.5, 3.5, 9.5])

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════

doc.save(DOC_PATH)
print("Done — gadash_report.docx updated successfully.")
