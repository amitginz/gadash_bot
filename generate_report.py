from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCR_BASE = r"C:\Users\amit ginzberg\Pictures\Screenshots"
SCR_DIR  = r"C:\Users\amit ginzberg\gadash_bot"

# Screenshots taken from live site
SCR_LOGIN         = SCR_DIR  + r"\scr_login.png"
SCR_MAIN_FULL     = SCR_DIR  + r"\scr_main.png"
SCR_MAIN_TABLE    = SCR_DIR  + r"\scr_main_table.png"
SCR_GLOBAL_SEARCH = SCR_DIR  + r"\scr_global_search.png"
SCR_QUICK_ADD     = SCR_DIR  + r"\scr_quick_add.png"
SCR_ADD           = SCR_DIR  + r"\scr_add.png"
SCR_SUMMARY       = SCR_DIR  + r"\scr_summary.png"
SCR_AUDIT         = SCR_DIR  + r"\scr_audit.png"
SCR_PRINT         = SCR_DIR  + r"\scr_print.png"
SCR_IMPORT        = SCR_DIR  + r"\scr_import.png"
SCR_API_DOCS      = SCR_DIR  + r"\scr_api_docs.png"
SCR_EDIT          = SCR_BASE + r"\Screenshot 2026-06-21 171017.png"
SCR_DELETE        = SCR_DIR  + r"\screen_delete.png"
SCR_LOGIN_WORKER  = SCR_DIR  + r"\scr_login_worker.png"
SCR_WORKER_LOGIN  = SCR_DIR  + r"\scr_worker_login.png"
SCR_WORKER_MAIN   = SCR_DIR  + r"\scr_worker_main.png"
SCR_FIELD_REPORT  = SCR_DIR  + r"\scr_field_report.png"
SCR_CLIENT_REPORT = SCR_DIR  + r"\scr_client_report.png"
SCR_FIELD_PRINT   = SCR_DIR  + r"\scr_field_print.png"

IMG_BASE = r"C:\Users\amit ginzberg\Documents\Amit_PC\Summaries\Computer_Science_Tel_Hai\Year_B\נושאים מתקדמים בפיתוח תוכנה\פרוייקט"
IMG_USECASE  = IMG_BASE + r"\UseCasedDiagram.png"
IMG_CLASS    = IMG_BASE + r"\ClassDiagram.png"
IMG_ACTIVITY = IMG_BASE + r"\ActivityDiagram.png"
IMG_STATE    = IMG_BASE + r"\StateMachineDiagram.png"

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── RTL helpers ───────────────────────────────────────────────────────────────
def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_rtl_para(doc, text, style='Normal', bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    p.alignment = align if align else WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    return p

def add_heading_rtl(doc, text, level=1):
    p = doc.add_heading('', level=level)
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_rtl_run(run)
    return p

def make_rtl_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_rtl_run(run)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A5F')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            set_rtl_run(run)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_picture_safe(doc, path, width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_rtl_para(doc, f'[תמונה לא נמצאה: {os.path.basename(path)}]',
                     size=9, color=(150, 150, 150), align=WD_ALIGN_PARAGRAPH.CENTER)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
add_rtl_para(doc, 'הפקולטה למדעים — החוג למדעי המחשב', bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_para(doc, 'נושאים מתקדמים בפיתוח תוכנה | סמסטר ב׳ תשפ״ו', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_para(doc, 'מרצה: ד״ר איאד סולימאן', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p_title)
r = p_title.add_run('מערכת ניהול עבודות גד״ש — אתר ובוט טלגרם')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
set_rtl_run(r)

add_rtl_para(doc, 'מסמך סופי', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_rtl_para(doc, 'כתובת המערכת:', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_para(doc, 'https://gadash-bot.fly.dev/', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_rtl_para(doc, 'שותפים:', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_para(doc, 'עמית גינזברג  |  ת״ז: 313393027  |  amitginz@gmail.com  |  054-3192907', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_rtl_para(doc, 'הדר קלר  |  ת״ז: 209825512  |  050-7551810', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_rtl_para(doc, 'תאריך הגשה: 15/08/2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — חלוקת עבודה
# ══════════════════════════════════════════════════════════════════════════════
add_heading_rtl(doc, 'סעיף 1 — חלוקת עבודה בין השותפים', 1)

make_rtl_table(doc,
    headers=['תחום', 'עמית גינזברג', 'הדר קלר'],
    rows=[
        ['ארכיטקטורת המערכת ו-WorkEntry Dataclass', '✓', ''],
        ['בוט טלגרם — ConversationHandler, מצבי שיחה, /undo, תמונות, סטטיסטיקות', '✓', ''],
        ['בוט — עריכת רשומה כלשהי, תזכורות לקוחות לא-פעילים', '✓', ''],
        ['שילוב Google Sheets (gspread, google-auth, cache)', '✓', ''],
        ['Flask Web App — נתיבים, לוגיקה עסקית, REST API, /api/docs', '✓', ''],
        ['אבטחה — CSRF, rate limiting, session timeout, סיסמה', '✓', ''],
        ['לוג שינויים (audit trail), דוח חודשי, דוח שבועי אוטומטי', '✓', ''],
        ['ממשק משתמש Web — HTML, CSS, Bootstrap RTL', '', '✓'],
        ['חיפוש גלובלי, הוספה מהירה (Quick-Add modal)', '', '✓'],
        ['גרפים בדוח חודשי (Chart.js × 3)', '', '✓'],
        ['עריכה מוטבעת (inline edit), מחיקה מרובה, שכפול רשומה', '', '✓'],
        ['מסך נייד (mobile cards), הדפסה/PDF', '', '✓'],
        ['טפסי הוספה/עריכה עם השלמה אוטומטית (autocomplete)', '', '✓'],
        ['ייצוא/ייבוא Excel וCSV, ייבוא חכם (dedup)', '', '✓'],
        ['פורטל עובד Web — /worker, WORKER_PASSWORD, היסטוריה אישית', '✓', ''],
        ['שדות גידול ושעות עבודה — WorkEntry, בוט, פורטל עובד וטפסי Web', '✓', ''],
        ['דוח שעות לפי חלקה/גידול (/field-report) — pivot + גרפים', '', '✓'],
        ['דוח לקוח (/client-report) — שעות, גרפים × 4, כל הרשומות', '', '✓'],
        ['הדפסת PDF של דוח חלקה (/field-report/print)', '', '✓'],
        ['תזכורת עובדים סוף-יום ב-18:00 + סיכום שבועי מורחב עם שעות', '✓', ''],
        ['פריסה לענן (Fly.io, Dockerfile, fly.toml)', '✓', ''],
        ['בדיקות אוטומטיות — pytest (27 בדיקות)', '✓', ''],
        ['תיעוד, README ו-UML', '✓', '✓'],
    ],
    col_widths=[9.5, 2.5, 2.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — תיאור מילולי
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_rtl(doc, 'סעיף 2 — תיאור מילולי של המערכת', 1)

add_heading_rtl(doc, 'רקע ומטרה', 2)
add_rtl_para(doc,
    'מערכת "ניהול עבודות גד״ש" פותחה עבור גוף חקלאי המנהל עבודות שדה (חריש, ריסוס, קציר, דיסוק). '
    'הצורך: עובדי שדה שאינם רגילים לממשקים ממוחשבים נדרשים לדווח על עבודות בזמן אמת, מהשדה; '
    'המנהל נדרש לצפות בנתונים מרוכזים, לסנן, לערוך ולייצא דוחות.',
    size=11)

add_heading_rtl(doc, 'שלושה ממשקים משלימים', 2)
add_rtl_para(doc,
    'בוט טלגרם — ממשק ניידי לעובד השדה: שיחה מודרכת בעברית (12 שלבי הזנה), '
    'תמיכה בתמונה כהערה, /undo לביטול הרשומה האחרונה, חיפוש וסטטיסטיקות. '
    'דוחות אוטומטיים: יומי ב-08:00, שבועי כל יום שני, תזכורת ב-18:00, '
    'והתרעה על לקוחות שלא עבדו ב-14 יום.',
    size=11)

add_rtl_para(doc,
    'פורטל עובד Web (/worker) — לעובד ללא טלגרם: טופס הזנה מהיר + היסטוריה אישית בלבד '
    '(ללא גישה לנתוני אחרים, עריכה, ייצוא או לוג). '
    'ממשק ניהול Web (/) — למנהל: טבלה עם עימוד, חיפוש גלובלי, סינון עם שמירה, '
    'Quick-Add modal, עריכה מוטבעת, מחיקה מרובה, שכפול, ייצוא/ייבוא Excel/CSV, '
    'דוח חודשי (3 גרפים), דוח חלקה, דוח לקוח, לוג שינויים, הדפסה/PDF ו-REST API.',
    size=11)

add_heading_rtl(doc, 'אחסון, אבטחה ופריסה', 2)
add_rtl_para(doc,
    'Google Sheets ("Gadash Data") הוא מקור האמת היחיד — נגיש ישירות למנהל, גיבוי אוטומטי, '
    'ללא עלות. cache בזיכרון (TTL 5 דקות) + threading.Lock למניעת race conditions. '
    'אבטחה: CSRF tokens, rate limiting (5 ניסיונות/דקה), session timeout 8 שעות, '
    'audit trail לכל פעולת כתיבה. '
    'פריסה: Fly.io (Amsterdam) — תהליך Python יחיד מריץ gunicorn ובוט ב-thread נפרד. '
    'כתובת: https://gadash-bot.fly.dev/',
    size=11)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — טכנולוגיות
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_rtl(doc, 'סעיף 3 — טכנולוגיות האינטרנט', 1)

add_heading_rtl(doc, '3.1 ארכיטקטורה', 2)
add_rtl_para(doc,
    'המערכת בנויה על ארכיטקטורה Monolithic בשכבות, הרצה בתהליך Python יחיד עם שני רכיבים '
    'במקביל: שרת Flask (gunicorn) לממשק ה-Web ובוט טלגרם ב-thread נפרד עם asyncio event loop. '
    'שני הרכיבים ניגשים לאותו מקור נתונים — Google Sheets — דרך ספריית gspread, '
    'עם מנגנון cache משותף ו-threading.Lock. '
    'נתוני הקלט עוברים אימות דרך WorkEntry dataclass לפני כל שמירה.',
    size=11)

add_heading_rtl(doc, '3.2 צד לקוח — Front-End', 2)
make_rtl_table(doc,
    headers=['טכנולוגיה', 'שימוש'],
    rows=[
        ['HTML5', 'מבנה כל דפי ה-Web'],
        ['CSS3', 'עיצוב מותאם, media queries למסך נייד'],
        ['Bootstrap 5.3 RTL', 'רספונסיביות, רכיבי UI, תמיכה בעברית ימין-לשמאל'],
        ['Bootstrap Icons', 'אייקונים (bi-*)'],
        ['JavaScript (Vanilla)', 'עריכה מוטבעת (double-click → PATCH), מחיקה מרובה, '
                                 'CSRF auto-inject, localStorage לשמירת סינון, '
                                 'השלמה אוטומטית (datalist)'],
        ['Chart.js', 'גרפי סטטיסטיקה בלוח הבקרה'],
        ['Fetch API', 'בקשות AJAX ל-REST API (PATCH /api/entries/<id>)'],
        ['Jinja2', 'תבניות HTML דינמיות (מנוע ה-templating של Flask)'],
        ['Telegram App', 'ממשק הנייד לעובדי השדה'],
    ],
    col_widths=[4.5, 10]
)

doc.add_paragraph()
add_heading_rtl(doc, '3.3 צד שרת — Back-End', 2)
make_rtl_table(doc,
    headers=['טכנולוגיה', 'גרסה', 'שימוש'],
    rows=[
        ['Python', '3.11', 'שפת התכנות הראשית'],
        ['Flask', '2.3.3', 'מסגרת Web — ניתוב HTTP, תבניות, session'],
        ['python-telegram-bot', '20.7', 'ספריית הבוט — ConversationHandler, async/await'],
        ['asyncio + threading', '—', 'הרצה מקבילה של הבוט ושרת ה-Web'],
        ['pandas', '2.2.2', 'עיבוד נתונים טבלאיים, קריאה/כתיבה Excel'],
        ['openpyxl', '3.1.2', 'מנוע ייצוא/ייבוא קבצי Excel'],
        ['gspread', '6.1.2', 'ממשק ל-Google Sheets API'],
        ['google-auth', '2.29.0', 'אימות מול Google (Service Account)'],
        ['gunicorn', '21.2.0', 'שרת WSGI לפרודקשן'],
        ['WorkEntry (dataclass)', '—', 'אימות ומבנה נתוני כל רשומת עבודה'],
    ],
    col_widths=[5, 2.5, 7]
)

doc.add_paragraph()
add_rtl_para(doc, 'נתיבי ה-Web המרכזיים:', bold=True, size=11)
make_rtl_table(doc,
    headers=['נתיב', 'Method', 'תיאור'],
    rows=[
        ['/', 'GET', 'דף ראשי — רשימת עבודות עם סינון, חיפוש ועימוד'],
        ['/login  /logout', 'GET, POST', 'כניסה מאוחדת (מנהל/עובד) עם rate limiting'],
        ['/add  /edit/<id>  /delete/<id>', 'GET, POST', 'הוספה, עריכה ומחיקת רשומה'],
        ['/bulk-delete  /duplicate/<id>', 'POST / GET', 'מחיקה מרובה ושכפול רשומה'],
        ['/summary  /audit  /print', 'GET', 'דוח חודשי, לוג שינויים, הדפסה/PDF'],
        ['/export  /export/csv  /import', 'GET / POST', 'ייצוא לאקסל/CSV וייבוא עם dedup'],
        ['/field-report  /client-report', 'GET', 'דוחות שעות לפי חלקה ולפי לקוח'],
        ['/worker', 'GET, POST', 'פורטל עובד — טופס הזנה + היסטוריה אישית'],
        ['/api/entries', 'GET / PATCH', 'REST API — רשומות JSON + עריכה מוטבעת'],
    ],
    col_widths=[4.5, 2.5, 7.5]
)

add_heading_rtl(doc, '3.4 בסיס נתונים', 2)
add_rtl_para(doc,
    'Google Sheets משמש כבסיס הנתונים של המערכת. הגיליון "Gadash Data" מכיל את העמודות הבאות:',
    size=11)
make_rtl_table(doc,
    headers=['שדה', 'תיאור'],
    rows=[
        ['שם לקוח', 'שם הלקוח / בית העסק (שדה חובה)'],
        ['תאריך', 'תאריך ביצוע (YYYY-MM-DD, שדה חובה)'],
        ['עבודה', 'סוג: חריש / ריסוס / קציר / דיסוק / אחר (שדה חובה)'],
        ['שם חלקה', 'שם או מזהה החלקה החקלאית'],
        ['גידול', 'סוג הגידול בחלקה (חיטה, תירס, כותנה... — רשות, autocomplete)'],
        ['כמות', 'כמות (למשל: "30 דונם")'],
        ['שעות', 'שעות עבודה (ערך מספרי, למשל: 3.5 — רשות)'],
        ['כלי', 'כלי העבודה (למשל: "טרקטור")'],
        ['מפעיל', 'שם מפעיל הציוד'],
        ['הערות', 'הערות חופשיות'],
        ['מזין', 'מי הזין — שם מטלגרם, "Web", או "import"'],
    ],
    col_widths=[4, 11]
)
doc.add_paragraph()
add_rtl_para(doc,
    'הסבר הבחירה ב-Google Sheets: נגיש ישירות לבעל העסק ללא ממשק נוסף, גיבוי אוטומטי '
    'ב-Google Drive, ניתן לשיתוף עם גורמים נוספים, ללא עלות, ו-API מובנה.',
    size=11)

add_heading_rtl(doc, '3.5 פלטפורמה', 2)
add_rtl_para(doc, 'המערכת מותאמת לשלוש פלטפורמות:', size=11)
add_rtl_para(doc, '• סמארטפון — ממשק הטלגרם, מותאם לשימוש בשדה', size=11)
add_rtl_para(doc, '• Web (דסקטופ/טאבלט) — ממשק הניהול עם טבלה, עריכה מוטבעת וגרפים', size=11)
add_rtl_para(doc, '• Web (נייד) — כרטיסיות mobile cards במקום טבלה ב-Bootstrap RTL', size=11)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — תדפיסי מסכים
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_rtl(doc, 'סעיף 4 — תדפיסי מסכים', 1)

screens = [
    ('מסך 1 — מסך כניסה מאוחד: בורר תפקיד (מנהל כחול / עובד ירוק)',    SCR_LOGIN),
    ('מסך 2 — לוח בקרה ראשי: טבלת רשומות, חיפוש גלובלי ועריכה מוטבעת', SCR_MAIN_TABLE),
    ('מסך 3 — Quick-Add modal: הוספת רשומה מהירה + טופס הוספה מלא',     SCR_QUICK_ADD),
    ('מסך 4 — פורטל עובד: טופס הזנה (גידול ושעות) + היסטוריה אישית',   SCR_WORKER_MAIN),
    ('מסך 5 — דוח שעות לפי חלקה וגידול: pivot table + גרפי בר ודונאט', SCR_FIELD_REPORT),
]
for title, path in screens:
    add_rtl_para(doc, title, bold=True, size=11)
    add_picture_safe(doc, path)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — UML
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_rtl(doc, 'סעיף 5 — מודלי UML', 1)
add_rtl_para(doc,
    'הדיאגרמות נוצרו ב-PlantUML ומצורפות כקבצים נפרדים (.puml). '
    'להלן תיאור של כל דיאגרמה.',
    size=11)

# 5.1 Use Case
add_heading_rtl(doc, '5.1 Use Case Diagram', 2)
add_picture_safe(doc, IMG_USECASE)
add_rtl_para(doc, 'שחקנים (Actors):', bold=True, size=11)
add_rtl_para(doc, '• עובד שדה — משתמש בבוט טלגרם להזנת עבודות וצפייה', size=11)
add_rtl_para(doc, '• עובד Web — משתמש בפורטל /worker להזנה וצפייה בהיסטוריה האישית', size=11)
add_rtl_para(doc, '• מנהל — משתמש בממשק ה-Web לניהול מלא', size=11)
add_rtl_para(doc, '• Google Sheets — מערכת חיצונית לאחסון', size=11)
doc.add_paragraph()
add_rtl_para(doc, 'תרחישי שימוש (Use Cases):', bold=True, size=11)
make_rtl_table(doc,
    headers=['Use Case', 'שחקן'],
    rows=[
        ['הזנת עבודה חדשה (שיחה מודרכת)', 'עובד שדה (בוט)'],
        ['צפייה באחרונות, חיפוש לפי לקוח, סטטיסטיקות', 'עובד שדה (בוט)'],
        ['ביטול הרשומה האחרונה (/undo)', 'עובד שדה (בוט)'],
        ['קבלת דוח יומי/שבועי אוטומטי ותזכורת 18:00', 'עובד שדה (בוט)'],
        ['הזנת עבודה + צפייה בהיסטוריה אישית', 'עובד Web (פורטל)'],
        ['צפייה בכל העבודות עם סינון, חיפוש ועימוד', 'מנהל'],
        ['הוספה (כולל Quick-Add), עריכה (כולל inline), שכפול, מחיקה', 'מנהל'],
        ['ייצוא Excel/CSV, ייבוא עם dedup', 'מנהל'],
        ['דוח חודשי (3 גרפים), דוח חלקה, דוח לקוח', 'מנהל'],
        ['לוג שינויים, הדפסה/PDF, שינוי סיסמה', 'מנהל'],
        ['גישה ל-REST API (JSON + inline PATCH)', 'מנהל / מפתח'],
        ['קריאה/כתיבה ב-Google Sheets', 'Google Sheets (מערכת)'],
    ],
    col_widths=[9, 5.5]
)

# 5.2 Class
doc.add_paragraph()
add_heading_rtl(doc, '5.2 Class Diagram', 2)
add_picture_safe(doc, IMG_CLASS)
add_rtl_para(doc, 'המערכת מחולקת ל-gadash/ package עם המודולים הבאים:', size=11)
make_rtl_table(doc,
    headers=['מודול / קובץ', 'תכונות / קבועים', 'פונקציות עיקריות'],
    rows=[
        ['gadash/models.py\nWorkEntry (dataclass)',
         'COLUMNS: list[str] — 11 עמודות\nVALID_TASKS: set\nclient, date, task, field_name\ncrop, amount, hours\ntool, operator, notes, entered_by',
         '__post_init__() — אימות חובה\nto_sheet_row(), to_dict()\nfrom_form(), from_bot(), from_dict()'],
        ['gadash/sheets.py\n(Google Sheets I/O)',
         '_gs_client: gspread.Client\n_cache_data, _cache_time\n_CACHE_TTL = 300 (שניות)\n_gs_lock: threading.Lock',
         'load_data_from_gsheet() → DataFrame\nappend_row_to_gsheet(entry)\nedit_row / delete_row / bulk_delete\npatch_cell_in_gsheet()\nsave_data_to_gsheet(df)'],
        ['gadash/bot.py\n(Telegram Bot)',
         'WEB_APP_URL: str\n16 מצבי שיחה (MENU…CONFIRM\nREGISTER_NAME, REGISTER_PASSWORD)\n_telegram_app, _telegram_loop',
         'start(), menu_choice(), confirm()\nnote_photo() — תמונה כהערה\nbot_undo(), bot_search_results()\n_scheduled_reports() — יומי/שבועי/18:00\nstart_telegram_bot()'],
        ['gadash/audit.py\n(Audit Log)',
         'AUDIT_LOG_FILE = "audit.log"\n_audit_queue: deque\n_audit_lock: Lock',
         '_log_audit(action, user, detail)\n_read_audit_log(limit) → list\n_flush_audit_to_sheets() — כל 30 שניות'],
        ['gadash/workers.py\n(Worker Accounts)',
         'SHA-256 password hashing',
         '_load_workers(), _verify_worker()\n_add_worker(), _delete_worker()\n_get_worker_by_telegram_id()\n_link_worker_telegram()'],
        ['app.py\n(Flask App)',
         'PAGE_SIZE = 50\n_current_password, _worker_password\nCSRF token + rate limiter',
         'index(), login(), add(), edit(), delete()\nbulk_delete(), duplicate(), summary()\naudit(), export(), import_data()\nfield_report(), client_report()\napi_entries(), api_entry_patch()'],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

# 5.3 Activity
doc.add_paragraph()
add_heading_rtl(doc, '5.3 Activity Diagram — זרימת הזנת עבודה דרך הבוט', 2)
add_picture_safe(doc, IMG_ACTIVITY)
add_rtl_para(doc, 'תיאור הזרימה:', bold=True, size=11)
for i, s in enumerate([
    'משתמש שולח /start',
    'הבוט מציג הודעת ברוך הבא ושואל אם להתחיל',
    'תפריט MENU — המשתמש בוחר: הזן עבודה / 5 אחרונות / חפש / סטטיסטיקות / ערוך אחרונה / סיים',
    'אם "הזן עבודה": הבוט מבקש שם לקוח (CLIENT)',
    'המשתמש מזין תאריך (YYYY-MM-DD או "היום") — אם שגוי, שואל שוב',
    'בחירת סוג עבודה מלוח מקשים (חריש / ריסוס / קציר / דיסוק / אחר)',
    'הזנת שם חלקה (FIELD)',
    'הזנת גידול (CROP) — ניתן לדלג עם כפתור "דלג"',
    'הזנת כמות (AMOUNT)',
    'הזנת שעות עבודה (HOURS) — ניתן לדלג עם כפתור "דלג"',
    'הזנת כלי, מפעיל, הערות (ניתן לדלג על הערות)',
    'WorkEntry.__post_init__() מאמת את כל השדות',
    'הצגת סיכום — המשתמש מאשר ("כן") או מבטל ("לא")',
    'אם אושר: append_row_to_gsheet() + _log_audit() ← חזרה לתפריט',
    'אם /undo: מחיקת הרשומה האחרונה של המשתמש + _log_audit("undo")',
    'אם "סיים": סיום השיחה',
], 1):
    add_rtl_para(doc, f'{i}. {s}', size=11)

# 5.4 State Machine
doc.add_paragraph()
add_heading_rtl(doc, '5.4 State Machine Diagram — מצבי שיחת הבוט', 2)
add_picture_safe(doc, IMG_STATE)
add_rtl_para(doc, 'מצבי המערכת (States):', bold=True, size=11)
make_rtl_table(doc,
    headers=['מצב', 'תיאור', 'מעברים אפשריים'],
    rows=[
        ['MENU', 'תפריט ראשי', 'CLIENT / SEARCH / [צפייה/סטטיסטיקות/עריכה] / END'],
        ['CLIENT', 'המתנה לשם לקוח', 'DATE'],
        ['DATE', 'המתנה לתאריך', 'TASK (תקין) / DATE (שגוי)'],
        ['TASK', 'בחירת סוג עבודה', 'FIELD'],
        ['FIELD', 'המתנה לשם חלקה', 'CROP'],
        ['CROP', 'המתנה לגידול (ניתן לדלג)', 'AMOUNT'],
        ['AMOUNT', 'המתנה לכמות', 'HOURS'],
        ['HOURS', 'המתנה לשעות עבודה (ניתן לדלג)', 'TOOL'],
        ['TOOL', 'המתנה לכלי', 'OPERATOR'],
        ['OPERATOR', 'המתנה למפעיל', 'NOTE'],
        ['NOTE', 'המתנה להערות (אפשר לדלג)', 'CONFIRM'],
        ['CONFIRM', 'אישור הנתונים', 'MENU (שמירה/ביטול)'],
        ['SEARCH', 'המתנה לשם לקוח לחיפוש', 'MENU (לאחר הצגת תוצאות)'],
        ['EDIT_SELECT', 'בחירת רשומה לצפייה (1-5)', 'MENU (לאחר הצגת פרטים וקישור)'],
        ['END', 'סיום השיחה', '—'],
    ],
    col_widths=[2.5, 4.5, 7.5]
)
doc.add_paragraph()
add_rtl_para(doc,
    'הערה: מכל מצב ניתן לשלוח /cancel (מעבר ל-END) או /start (מעבר ל-MENU). '
    '/undo זמין כ-fallback מכל מצב.',
    size=10)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — GitHub
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_rtl(doc, 'סעיף 6 — ניהול גרסאות ב-GitHub (בונוס)', 1)
add_rtl_para(doc,
    'הפרויקט מנוהל ב-Git עם מאגר ב-GitHub תחת חשבון amitginz. '
    'כתובת המאגר: https://github.com/amitginz/gadash_bot',
    size=11)
doc.add_paragraph()
make_rtl_table(doc,
    headers=['Commit Hash', 'תיאור'],
    rows=[
        ['ebfd1be', 'Add work hours & crop tracking, field/client reports, PDF print, worker reminders'],
        ['0774f0a', 'Unify login: single /login page for both manager and worker with role selector'],
        ['efaef93', 'Add worker portal: separate /worker page with personal history'],
        ['1072140', 'Add global search, quick-add modal, charts, API docs, pytest (27 tests)'],
        ['1caf955', 'Add CSRF, rate limit, session timeout, inline edit, audit log, REST API'],
        ['67b8f9f', 'Refactor: add WorkEntry dataclass for validated data management'],
        ['0150cb1', 'Deploy to Fly.io — Dockerfile, fly.toml, gunicorn'],
    ],
    col_widths=[3.5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ספריות צד שלישי
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_rtl(doc, 'סעיף 7 — ספריות צד שלישי', 1)
make_rtl_table(doc,
    headers=['ספרייה', 'גרסה', 'שימוש'],
    rows=[
        ['Flask', '2.3.3', 'מסגרת Web, ניתוב, תבניות Jinja2, session management'],
        ['python-telegram-bot', '20.7', 'בוט טלגרם — async, ConversationHandler, JobQueue'],
        ['pandas', '2.2.2', 'עיבוד נתונים טבלאיים, Excel, pivot tables'],
        ['numpy', '1.26.4', 'תלות של pandas'],
        ['openpyxl', '3.1.2', 'קריאה/כתיבה קבצי Excel (.xlsx)'],
        ['gspread', '6.1.2', 'גישה ל-Google Sheets API'],
        ['google-auth', '2.29.0', 'אימות מול Google Service Account'],
        ['gunicorn', '21.2.0', 'שרת WSGI לפרודקשן'],
        ['Bootstrap', '5.3.0 RTL', 'עיצוב UI, רספונסיביות, תמיכה עברית'],
        ['Bootstrap Icons', '1.11.0', 'אייקוני SVG (bi-*)'],
        ['Chart.js', '4.x (CDN)', 'גרפי סטטיסטיקה בדפדפן (עוגה, עמודות, אופקי)'],
        ['pytest', '9.x', 'בדיקות אוטומטיות — 27 בדיקות ל-WorkEntry ולנתיבי Flask'],
    ],
    col_widths=[4.5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = 'gadash_report.docx'
doc.save(out)
print(f'Saved: {out}')
